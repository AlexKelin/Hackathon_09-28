import os
from flask import Flask, request, render_template, jsonify
from werkzeug.utils import secure_filename
import requests
from docx import Document
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_caching import Cache
from dotenv import load_dotenv
import json
from ai21 import AI21Client
from ai21.models.chat import ChatMessage, ToolMessage
from ai21.models.chat.function_tool_definition import FunctionToolDefinition
from ai21.models.chat.tool_defintions import ToolDefinition
from ai21.models.chat.tool_parameters import ToolParameters
import logging
from flask_cors import CORS
import re
import random  # Add this import at the top of the file

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)  # This will enable CORS for all routes
app.config['DEBUG'] = True  # Enable debug mode

# Configuration
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload size
app.config['ALLOWED_EXTENSIONS'] = {'docx'}
app.config['MAX_FILES'] = 3
app.config['JAMBO_AI_API_URL'] = os.getenv('JAMBO_AI_API_URL')
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY')

# Setup rate limiting
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://"
)

# Setup caching
cache = Cache(app, config={'CACHE_TYPE': 'simple'})

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/', methods=['GET', 'POST'])
@limiter.limit("10 per minute")
def upload_files():
    print("Starting script...")
    if request.method == 'POST':
        try:
            if 'files[]' not in request.files:
                return jsonify({'error': 'No resume files uploaded'}), 400
            
            files = request.files.getlist('files[]')
            
            # Handle job description
            job_description = request.form.get('job_description', '')
            job_description_file = request.files.get('job_description_file')
            
            if job_description_file and allowed_file(job_description_file.filename):
                filename = secure_filename(job_description_file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                job_description_file.save(file_path)
                
                # Read the job description from the file
                if filename.endswith('.docx'):
                    doc = Document(file_path)
                    job_description = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                else:  # Assume it's a .txt file
                    with open(file_path, 'r') as file:
                        job_description = file.read()
                
                # Clean up the temporary file
                os.remove(file_path)
            
            if not job_description:
                return jsonify({'error': 'No job description provided'}), 400
            
            if len(files) > app.config['MAX_FILES']:
                return jsonify({'error': f"Maximum {app.config['MAX_FILES']} resume files allowed"}), 400
            
            uploaded_files = []
            for file in files:
                if file and allowed_file(file.filename):
                    filename = secure_filename(file.filename)
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    file.save(file_path)
                    uploaded_files.append(file_path)
            
            if uploaded_files:
                app.logger.info("Processing files...")
                result = process_files(uploaded_files, job_description)
                app.logger.info(f"Result: {result}")
                return jsonify(result)
            
            return jsonify({'error': 'No valid resume files uploaded'}), 400
        
        except Exception as e:
            app.logger.exception("An error occurred during file processing")
            return jsonify({'error': str(e)}), 500
    
    return render_template('upload.html')

@cache.memoize(timeout=300)  # Cache for 5 minutes
def process_files(file_paths_or_texts, job_description):
    print("Starting main logic...")
    documents_text = []
    for item in file_paths_or_texts:
        if isinstance(item, str) and os.path.isfile(item):
            # It's a file path
            doc = Document(item)
            full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            # It's already text
            full_text = item
        documents_text.append(full_text)
    
    # Initialize AI21Client
    AI21_API_KEY = os.environ.get('AI21_API_KEY')
    client = AI21Client(api_key=AI21_API_KEY)

    # Define the tool for resume screening
    tool_definition = ToolDefinition(
        type="function",
        function=FunctionToolDefinition(
            name="screen_resumes",
            description="Screen resumes against a job description",
            parameters=ToolParameters(
                type="object",
                properties={
                    "resumes": {"type": "array", "items": {"type": "string"}, "description": "List of resume texts"},
                    "job_description": {"type": "string", "description": "Job description text"}
                },
                required=["resumes", "job_description"],
            ),
        ),
    )

    # Create messages for the chat
    messages = [
        ChatMessage(
            role="system",
            content="You are a helpful resume screening assistant. Use the supplied tool to screen resumes against a job description."
        ),
        ChatMessage(role="user", content=f"Please screen these resumes against the following job description: {job_description}")
    ]

    # Make the initial request to the model
    response = client.chat.completions.create(messages=messages, model="jamba-1.5-large", tools=[tool_definition])

    # Process the response and handle tool calls
    assistant_message = response.choices[0].message
    messages.append(assistant_message)
    result = None
    tool_calls = assistant_message.tool_calls

    if tool_calls:
        tool_call = tool_calls[0]
        if tool_call.function.name == "screen_resumes":
            func_arguments = json.loads(tool_call.function.arguments)
            if "resumes" in func_arguments and "job_description" in func_arguments:
                # Implement actual resume screening logic here
                screened_resumes = []
                for resume in documents_text:
                    score, feedback, division_lead_thought, hr_lead_thought, ceo_thought, hiring_committee_decision_result = screen_resume(resume, job_description)
                    screened_resumes.append({
                        "score": score,
                        "feedback": feedback,
                        "division_lead_thought": division_lead_thought,
                        "hr_lead_thought": hr_lead_thought,
                        "ceo_thought": ceo_thought,
                        "hiring_committee_decision": hiring_committee_decision_result
                    })
                result = {"screened_resumes": screened_resumes}
            else:
                app.logger.error("Missing 'resumes' or 'job_description' in function arguments")
        else:
            app.logger.error(f"Unexpected tool call found - {tool_call.function.name}")
    else:
        app.logger.error("No tool calls found")

    if result is not None:
        tool_message = ToolMessage(role="tool", tool_call_id=tool_calls[0].id, content=json.dumps(result))
        messages.append(tool_message)
        response = client.chat.completions.create(messages=messages, model="jamba-1.5-large", tools=[tool_definition])
        
        # Extract and return the content
        return {
            "api_response": response.choices[0].message.content,
            "screened_resumes": result["screened_resumes"]
        }
    else:
        return {'error': 'Failed to process applicant screening'}

def ceo_evaluation(resume, job_description):
    # Initialize AI21Client
    AI21_API_KEY = os.environ.get('AI21_API_KEY')
    client = AI21Client(api_key=AI21_API_KEY)

    messages = [
        ChatMessage(
            role="system",
            content="You are the CEO of the company. Briefly evaluate if this candidate would be a good cultural fit based on their resume and the job description. Provide a short, one-sentence response."
        ),
        ChatMessage(
            role="user",
            content=f"Job Description:\n{job_description}\n\nResume:\n{resume}\n\nAs the CEO, what's your brief thought on this candidate's cultural fit?"
        )
    ]

    response = client.chat.completions.create(messages=messages, model="jamba-1.5-large", max_tokens=50)
    return response.choices[0].message.content.strip()

def division_lead_evaluation(resume, job_description):
    # Initialize AI21Client
    AI21_API_KEY = os.environ.get('AI21_API_KEY')
    client = AI21Client(api_key=AI21_API_KEY)

    messages = [
        ChatMessage(
            role="system",
            content="You are the lead of the division. Briefly evaluate if this candidate would be a good fit for your team based on their resume and the job description. Provide a short, one-sentence response."
        ),
        ChatMessage(
            role="user",
            content=f"Job Description:\n{job_description}\n\nResume:\n{resume}\n\nAs the division lead, what's your brief thought on this candidate's fit for your team?"
        )
    ]

    response = client.chat.completions.create(messages=messages, model="jamba-1.5-large", max_tokens=50)
    return response.choices[0].message.content.strip()

def hr_lead_evaluation(resume, job_description):
    # Initialize AI21Client
    AI21_API_KEY = os.environ.get('AI21_API_KEY')
    client = AI21Client(api_key=AI21_API_KEY)

    messages = [
        ChatMessage(
            role="system",
            content="You are the HR lead. Briefly evaluate if this candidate would be a good fit for the company culture and values based on their resume and the job description. Provide a short, one-sentence response."
        ),
        ChatMessage(
            role="user",
            content=f"Job Description:\n{job_description}\n\nResume:\n{resume}\n\nAs the HR lead, what's your brief thought on this candidate's fit for the company culture and values?"
        )
    ]

    response = client.chat.completions.create(messages=messages, model="jamba-1.5-large", max_tokens=50)
    return response.choices[0].message.content.strip()

def hiring_committee_decision(resume, job_description, division_lead_thought, hr_lead_thought, ceo_thought, score, feedback):
    # Initialize AI21Client
    AI21_API_KEY = os.environ.get('AI21_API_KEY')
    client = AI21Client(api_key=AI21_API_KEY)

    messages = [
        ChatMessage(
            role="system",
            content="You are the Hiring Committee. Based on all the accumulated feedback, make a final decision on whether to approve or not approve the candidate. Provide a short, one-sentence response starting with either 'Approves:' or 'Does Not Approve:'."
        ),
        ChatMessage(
            role="user",
            content=f"Job Description:\n{job_description}\n\nResume:\n{resume}\n\nDivision Lead's Thought: {division_lead_thought}\nHR Lead's Thought: {hr_lead_thought}\nCEO's Thought: {ceo_thought}\nScore: {score}\nFeedback: {feedback}\n\nAs the Hiring Committee, what's your final decision on this candidate?"
        )
    ]

    response = client.chat.completions.create(messages=messages, model="jamba-1.5-large", max_tokens=50)
    return response.choices[0].message.content.strip()

def screen_resume(resume, job_description):
    # Initialize AI21Client
    AI21_API_KEY = os.environ.get('AI21_API_KEY')
    client = AI21Client(api_key=AI21_API_KEY)

    # Create messages for the resume screening
    messages = [
        ChatMessage(
            role="system",
            content="You are an expert resume screener. Evaluate the resume against the job description and provide a score and detailed feedback."
        ),
        ChatMessage(
            role="user",
            content=f"Job Description:\n{job_description}\n\nResume:\n{resume}\n\nPlease evaluate this resume against the job description. Provide a score from 0 to 1 and detailed feedback on the match, including specific skills, experience, and qualifications that align or are missing."
        )
    ]

    # Make the request to the model
    response = client.chat.completions.create(messages=messages, model="jamba-1.5-large", max_tokens=500)

    # Extract the score and feedback from the response
    content = response.choices[0].message.content
    score_match = re.search(r'Score:\s*(0\.\d+|1\.0)', content)
    score = float(score_match.group(1)) if score_match else 0.0
    feedback = content.split("Feedback:", 1)[-1].strip()

    # Get division lead's evaluation
    division_lead_thought = division_lead_evaluation(resume, job_description)

    # Get HR lead's evaluation
    hr_lead_thought = hr_lead_evaluation(resume, job_description)

    # Get CEO's evaluation
    ceo_thought = ceo_evaluation(resume, job_description)

    # Get Hiring Committee's final decision
    hiring_committee_decision_result = hiring_committee_decision(resume, job_description, division_lead_thought, hr_lead_thought, ceo_thought, score, feedback)

    return score, feedback, division_lead_thought, hr_lead_thought, ceo_thought, hiring_committee_decision_result

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({'error': 'File too large'}), 413

@app.errorhandler(429)
def ratelimit_handler(e):
    return jsonify({'error': "Rate limit exceeded"}), 429

@app.route('/test_api', methods=['POST'])
def test_api():
    data = request.json
    job_description = data.get('job_description', '')
    resume = data.get('resume', '')
    
    client = AI21Client(api_key=AI21_API_KEY)
    
    messages = [
        ChatMessage(
            role="system",
            content="You are a helpful resume screening assistant. Evaluate the resume against the job description and provide a score and feedback."
        ),
        ChatMessage(
            role="user",
            content=f"Job Description: {job_description}\n\nResume: {resume}\n\nPlease evaluate this resume against the job description."
        )
    ]

    try:
        response = client.chat.completions.create(
            messages=messages,
            model="jamba-1.5-large",
            max_tokens=300
        )
        return jsonify({"response": response.choices[0].message.content})
    except Exception as e:
        logger.exception("An error occurred during API test")
        return jsonify({'error': f'An error occurred: {str(e)}'}), 500

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    app.run(host='0.0.0.0', port=5001)  # Change to a different port, e.g., 5001